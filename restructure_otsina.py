from docx import Document

doc = Document()

# Page de garde
doc.add_paragraph("MINISTERE DE L'ECONOMIE NUMERIQUE\tREPUBLIQUE GABONAISE")
doc.add_paragraph("Union – Travail – Justice")
doc.add_paragraph("")
doc.add_paragraph("INSTITUT NATIONAL DE LA POSTE, DES TECHNOLOGIES DE L'INFORMATION ET DE LA COMMUNICATION")
doc.add_paragraph("")
doc.add_paragraph("RAPPORT DE STAGE DE FIN DE FORMATION POUR L'OBTENTION DU")
doc.add_paragraph("DIPLOME DE LICENCE PROFESSIONNELLE")
doc.add_paragraph("")
doc.add_paragraph("THEME : Développement d'un système informatisé de facturation et de suivi de stock")
doc.add_paragraph("")
doc.add_paragraph("OPTION : DÉVELOPPEMENT D'APPLICATIONS RÉPARTIES")
doc.add_paragraph("")
doc.add_paragraph("Présenté et soutenu par :")
doc.add_paragraph("OTSINA MINKO Jean Rodrigue Rismin")
doc.add_paragraph("")
doc.add_paragraph("Sous la direction de :")
doc.add_paragraph("Tuteur en entreprise : M. SAMBA SONKO")

doc.add_page_break()

# Dédicaces
doc.add_heading("DEDICACES", level=1)
doc.add_paragraph("Je dédie ce rapport à ma famille, mes proches et toutes les personnes qui m'ont soutenu tout au long de mon parcours. Leur confiance, leurs encouragements et leur présence ont été des piliers essentiels dans la réussite de cette formation et de ce stage.")

# Remerciements
doc.add_heading("REMERCIEMENTS", level=1)
doc.add_paragraph("Je tiens à exprimer mes sincères remerciements à TECH INFO PLUS, représentée par son Directeur Général Monsieur SONKO SAMBA MONSESILO, pour m'avoir accueilli en stage et offert l'opportunité de travailler sur un projet stratégique pour l'entreprise.")
doc.add_paragraph("Ma gratitude s'adresse tout particulièrement à mon maître de stage, Monsieur SONKO SAMBA MONSESILO, pour son encadrement, ses conseils techniques et méthodologiques, ainsi que pour la confiance accordée dans la réalisation de ce système de gestion commerciale.")
doc.add_paragraph("Je remercie également l'ensemble des collaborateurs de TECH INFO PLUS pour leur disponibilité, leurs retours constructifs et leur collaboration durant toutes les phases du projet. Leur implication a été déterminante pour la qualité du produit final.")
doc.add_paragraph("Enfin, mes remerciements vont à mon encadrant académique, Monsieur ENGOANG DALADIER, pour son suivi régulier et ses recommandations, ainsi qu'à ma famille pour son soutien moral constant.")

# Résumé
doc.add_heading("RESUME", level=1)
doc.add_paragraph("Ce rapport retrace le stage de fin de formation effectué au sein de TECH INFO PLUS, une entreprise spécialisée dans les solutions informatiques et les services de télécommunication. La mission principale était de concevoir et développer un système informatisé de facturation et de suivi de stock afin de moderniser les processus de gestion commerciale.")
doc.add_paragraph("La solution livrée est une plateforme web basée sur une architecture 3-tiers : une interface utilisateur développée avec React.js, une API backend construite avec FastAPI, et une base de données relationnelle hébergée sur Supabase/PostgreSQL. Le système implémente quinze modules fonctionnels couvrant l'ensemble du cycle de gestion (clients, articles, devis, facturation, ventes comptoir, stock, règlements, rapports, etc.). L'authentification repose sur JWT et la sécurité des mots de passe sur bcrypt.")
doc.add_paragraph("L'automatisation permet de réduire 85 % des processus manuels, de diminuer de plus de 95 % les erreurs et d'accroître la productivité globale : la génération des factures est passée de dix minutes à deux minutes, les rapports sont produits en quelques secondes, et la traçabilité est désormais complète.")
doc.add_paragraph("Au plan pédagogique, ce stage a permis de renforcer les compétences en développement web full-stack, en gestion de projet Agile/Scrum et en collaboration avec des parties prenantes non techniques.")

# Abstract
doc.add_heading("ABSTRACT", level=1)
doc.add_paragraph("This report presents the professional internship carried out at TECH INFO PLUS, a company specialized in IT solutions and telecommunication services. The primary objective was to design and implement an integrated web-based system for invoicing and inventory tracking, in order to streamline the company's commercial operations.")
doc.add_paragraph("The delivered solution relies on a three-tier architecture combining a React.js frontend, a FastAPI backend and a relational database hosted on Supabase/PostgreSQL. Fifteen business modules were implemented to cover the entire commercial lifecycle. JWT authentication and bcrypt password hashing ensure strong security.")
doc.add_paragraph("The system automates 85% of manual processes, reduces errors by 95%, and boosts productivity. Invoice issuance time dropped from ten minutes to two, reports can be generated within seconds, and data traceability is guaranteed.")
doc.add_paragraph("From a learning perspective, the internship strengthened full-stack development skills, agile project management abilities, and collaboration with non-technical stakeholders.")

# Introduction
doc.add_heading("INTRODUCTION", level=1)
doc.add_paragraph("La transformation numérique des processus de gestion représente aujourd'hui un enjeu majeur pour les entreprises gabonaises, qui doivent concilier performance opérationnelle, fiabilité des données et satisfaction des clients. TECH INFO PLUS, entreprise basée à Port-Gentil et spécialisée dans les solutions informatiques, a souhaité moderniser ses outils internes afin de fiabiliser ses opérations commerciales et de soutenir sa croissance.")
doc.add_paragraph("Le stage réalisé au sein du service technique avait pour objectif de concevoir et développer un système informatisé de facturation et de suivi de stock répondant précisément aux besoins de l'entreprise. Le projet s'est inscrit dans une démarche méthodique, articulée autour de l'analyse des pratiques existantes, de la définition détaillée des besoins, de la conception d'une architecture logicielle robuste et de la mise en œuvre progressive des fonctionnalités.")
doc.add_paragraph("Ce rapport est structuré en deux parties. La première présente l'environnement professionnel du stage et le cadre théorique sur lequel le projet s'est appuyé. La seconde détaille le déroulement du stage, les réalisations techniques et les solutions apportées à la problématique initiale, avant de conclure par une synthèse des résultats obtenus et des compétences développées.")

# PARTIE I
doc.add_heading("PARTIE I : PRESENTATION DE LA STRUCTURE D'ACCUEIL", level=1)

doc.add_heading("I - Présentation de la structure d'accueil", level=2)

doc.add_heading("1. Historique et positionnement", level=3)
doc.add_paragraph("TECH INFO PLUS est une société gabonaise fondée en 2009 par Monsieur SONKO SAMBA MONSESILO. Implantée au cœur de Port-Gentil, elle intervient dans le domaine des technologies de l'information et des télécommunications. L'entreprise propose des prestations de vidéo surveillance, de liaisons spécialisées, de gestion de réseaux, de maintenance informatique et de développement de solutions logicielles sur mesure. Au fil des années, TECH INFO PLUS s'est imposée comme un partenaire de confiance pour les PME locales grâce à son expertise technique et à sa capacité d'adaptation aux besoins spécifiques de ses clients.")

doc.add_heading("2. Organisation et gouvernance", level=3)
doc.add_paragraph("La structure organisationnelle de TECH INFO PLUS s'articule autour de quatre pôles complémentaires : le service de maintenance informatique, le service réseaux et télécommunications, le service infographie, et le service développement. Le Directeur Général assume également la fonction de chef de projet pour les développements logiciels et veille à la coordination entre les équipes techniques et les clients.")

doc.add_heading("3. Missions et services offerts", level=3)
doc.add_paragraph("Les prestations de TECH INFO PLUS couvrent un large spectre : installation et maintenance de parcs informatiques, développement d'applications web, mise en place d'infrastructures réseau, formation des utilisateurs et fourniture d'équipements informatiques. L'entreprise se démarque par sa capacité à proposer des solutions intégrées associant matériel, logiciels et accompagnement. Elle intervient autant auprès des entreprises qu'auprès des particuliers exigeant un haut niveau de qualité de service.")

doc.add_heading("4. Ressources humaines et matérielles", level=3)
doc.add_paragraph("L'équipe technique regroupe des profils expérimentés en développement, en administration système et réseau, ainsi qu'en support utilisateur. Les collaborateurs s'appuient sur un parc matériel complet et sur une stack logicielle professionnelle comprenant Visual Studio Code, Git, Node.js, Python, suites bureautiques et outils de supervision réseau.")

doc.add_heading("II - Cadre théorique", level=2)

doc.add_heading("1. Cahier des charges du projet", level=3)
doc.add_paragraph("Le projet de stage visait la conception et la réalisation d'un système de gestion commerciale permettant d'automatiser la facturation, de sécuriser les flux financiers et de suivre rigoureusement les stocks. Les objectifs spécifiques portaient sur la numérisation des documents commerciaux, la traçabilité des opérations, la génération de rapports fiables et la mise en place d'un contrôle d'accès par rôles.")

doc.add_heading("2. Analyse critique de l'existant", level=3)
doc.add_paragraph("Avant le projet, la gestion s'appuyait sur des outils bureautiques dispersés tels que des cahiers, des feuilles Excel et des documents Word. Cette organisation générait de nombreuses erreurs de calcul, une perte de temps importante, des ruptures de stock mal anticipées et une absence de visibilité pour la direction. Les employés avaient des difficultés à consolider les données et à éditer des rapports fiables.")

doc.add_heading("3. Solutions envisagées et choix retenus", level=3)
doc.add_paragraph("Trois options ont été étudiées : adopter une solution SaaS existante, personnaliser une solution open-source ou développer un système sur mesure. La troisième option a été retenue car elle garantissait une adéquation totale aux besoins métiers, la maîtrise des données, la flexibilité pour incorporer des évolutions futures et un coût maîtrisé sur le long terme.")

doc.add_heading("4. Choix architecturaux et technologiques", level=3)
doc.add_paragraph("La solution repose sur une architecture 3-tiers. Le frontend est développé en React.js afin d'offrir une interface réactive et maintenable. Le backend est construit avec FastAPI pour profiter de performances élevées et d'une documentation automatique. La base de données relationnelle est hébergée sur Supabase/PostgreSQL, avec SQLAlchemy pour l'abstraction des requêtes. L'authentification s'appuie sur JWT et la sécurisation des mots de passe sur bcrypt.")

doc.add_heading("5. Référentiel méthodologique", level=3)
doc.add_paragraph("Le projet a été mené selon la méthodologie Agile/Scrum, structurée en six sprints de deux semaines. Chaque itération comprenait une phase de planification, de développement, de tests et de revue avec les utilisateurs. Cette approche a favorisé l'incrémentalité, la prise en compte rapide des retours et la sécurisation des livrables.")

# PARTIE II
doc.add_heading("PARTIE II : TRAITEMENT DU THEME", level=1)

doc.add_heading("I - Déroulement du stage", level=2)

doc.add_heading("1. Phase d'intégration et immersion terrain", level=3)
doc.add_paragraph("Le stage a débuté par une phase d'intégration permettant de comprendre l'organisation interne, de rencontrer les différentes équipes et de recenser les besoins. Des entretiens ont été menés avec la direction, le service commercial et les agents opérationnels afin d'identifier les irritants du système existant et de formaliser les attentes fonctionnelles.")

doc.add_heading("2. Planification et organisation des sprints", level=3)
doc.add_paragraph("Un backlog fonctionnel a été constitué à partir du cahier des charges. Les fonctionnalités ont été priorisées pour sécuriser rapidement les éléments critiques tels que l'authentification, la facturation et la gestion des clients. Chaque sprint de deux semaines visait la livraison d'un incrément testable, accompagné d'une démonstration et d'une revue. Les itérations ont couvert successivement la modélisation, l'infrastructure, le MVP, les modules avancés, les tests et le déploiement.")

doc.add_heading("3. Mise en œuvre des solutions techniques", level=3)
doc.add_paragraph("Le développement s'est appuyé sur une séparation stricte entre le frontend et le backend. L'API FastAPI expose trente-deux endpoints sécurisés, structurés autour des ressources principales telles que les utilisateurs, les clients, les articles, les devis, les factures, les règlements et les mouvements de stock. Le frontend React.js consomme ces endpoints via Axios et gère l'état global au moyen de hooks dédiés. La base de données PostgreSQL a été modélisée avec quinze tables principales et des clés étrangères assurant l'intégrité référentielle.")

doc.add_heading("4. Collaboration et gestion des tests", level=3)
doc.add_paragraph("Des sessions de tests fonctionnels ont été réalisées avec les futurs utilisateurs pour valider chaque module. Les retours ont conduit à plusieurs itérations correctives portant sur l'ergonomie des formulaires, la validation des seuils de stock et les filtres des rapports. Des scripts de tests API ont été exécutés via Postman pour garantir la conformité aux exigences, et une procédure de recette a été établie avant le déploiement définitif.")

doc.add_heading("II - Résolution de la problématique", level=2)

doc.add_heading("1. Architecture fonctionnelle livrée", level=3)
doc.add_paragraph("La solution comprend quinze modules interconnectés couvrant l'ensemble du cycle commercial : gestion des utilisateurs, gestion des clients, gestion des articles, fournisseurs, devis, facturation (trois types), ventes comptoir, règlements, avoirs, gestion du stock, inventaires, rapports statistiques, configuration du système et suivi des anomalies. Chaque module applique une nomenclature automatique des documents et assure la traçabilité des opérations.")

doc.add_heading("2. Sécurité, performance et traçabilité", level=3)
doc.add_paragraph("L'accès au système est conditionné par une authentification JWT et un contrôle d'accès basé sur les rôles (administrateur, vendeur, comptable). Les mots de passe sont hachés avec bcrypt. Le backend intègre des middlewares CORS et une journalisation des requêtes. Les opérations sensibles, telles que la création de facture ou l'ajustement de stock, sont historisées afin de retracer l'auteur, la date et le détail des actions.")

doc.add_heading("3. Résultats obtenus pour TECH INFO PLUS", level=3)
doc.add_paragraph("Les indicateurs clés démontrent l'impact de la solution : réduction de 80 % du temps de facturation, diminution de 95 % des erreurs de calcul, génération quasi instantanée des rapports et gain substantiel sur le suivi des stocks. Les responsables disposent désormais d'un tableau de bord consolidé et peuvent piloter l'activité à partir de données fiables.")

doc.add_heading("4. Difficultés rencontrées et leçons tirées", level=3)
doc.add_paragraph("Plusieurs défis ont été relevés : synchronisation de l'état dans React, gestion des relations complexes entre les tables, génération de PDF professionnels et migration de MySQL vers PostgreSQL. Ces difficultés ont été surmontées par une veille technique continue, la mutualisation des ressources documentaires et la mise en place de scripts de migration contrôlés.")

# Conclusion
doc.add_heading("CONCLUSION", level=1)
doc.add_paragraph("Le stage réalisé au sein de TECH INFO PLUS a permis de livrer une solution complète d'automatisation de la facturation et de suivi de stock, répondant aux attentes de l'entreprise et améliorant significativement son efficacité opérationnelle. Le système mis en production est aujourd'hui utilisé quotidiennement par les collaborateurs et constitue un levier majeur de modernisation des processus internes.")
doc.add_paragraph("Sur le plan personnel, cette expérience a consolidé les compétences en développement web full-stack, en conception d'architectures logicielles sécurisées et en gestion de projet agile. Elle a également renforcé l'aptitude à communiquer avec des utilisateurs non techniques et à transformer leurs besoins en solutions concrètes.")
doc.add_paragraph("Les perspectives d'évolution incluent l'intégration d'une application mobile pour les commerciaux, l'automatisation des envois de factures par email, la connexion à des solutions comptables externes et, à terme, la transformation du système en plateforme SaaS multi-entreprises.")

# Bibliographie
doc.add_heading("BIBLIOGRAPHIE / WEBOGRAPHIE", level=1)
biblio = [
    "Documentation React.js – https://react.dev/",
    "Documentation FastAPI – https://fastapi.tiangolo.com/",
    "Documentation PostgreSQL – https://www.postgresql.org/docs/",
    "Documentation SQLAlchemy – https://docs.sqlalchemy.org/",
    "Documentation Chart.js – https://www.chartjs.org/",
    "Documentation jsPDF – https://github.com/parallax/jsPDF",
    "Object Management Group (UML) – https://www.omg.org/spec/UML/",
    "W3Schools – https://www.w3schools.com/",
    "MDN Web Docs – https://developer.mozilla.org/",
    "Stack Overflow – https://stackoverflow.com/",
    "GitHub – https://github.com/",
    "Robert C. Martin (2008). Clean Code: A Handbook of Agile Software Craftsmanship.",
    "Erich Gamma et al. (1994). Design Patterns: Elements of Reusable Object-Oriented Software.",
    "Robert C. Martin (2002). Agile Software Development, Principles, Patterns, and Practices."
]
for item in biblio:
    doc.add_paragraph(item)

# Annexes
doc.add_heading("ANNEXES", level=1)
annexes = [
    "Annexe 1 : Diagrammes UML",
    "Annexe 2 : Captures d'écran de l'application",
    "Annexe 3 : Exemples de documents générés",
    "Annexe 4 : Schéma de la base de données",
    "Annexe 5 : Architecture technique et déploiement",
    "Annexe 6 : Planning du projet"
]
for annexe in annexes:
    doc.add_paragraph(annexe)

doc.save("OTSINA TAYLOR.docx")
print("✅ Document OTSINA TAYLOR.docx restructuré avec succès !")




